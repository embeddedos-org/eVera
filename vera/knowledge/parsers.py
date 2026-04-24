"""Document parsers for the knowledge base.

Supports PDF, DOCX, TXT, Markdown, and CSV formats.
"""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_pdf(content: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(content))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n\n".join(texts)
    except ImportError:
        logger.warning("PyPDF2 not installed — cannot parse PDFs")
        return ""
    except Exception as e:
        logger.error("Failed to parse PDF: %s", e)
        return ""


def parse_docx(content: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        texts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                texts.append(paragraph.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texts.append(row_text)
        return "\n\n".join(texts)
    except ImportError:
        logger.warning("python-docx not installed — cannot parse DOCX")
        return ""
    except Exception as e:
        logger.error("Failed to parse DOCX: %s", e)
        return ""


def parse_txt(content: bytes) -> str:
    """Extract text from a plain text or Markdown file."""
    try:
        import chardet

        detected = chardet.detect(content)
        encoding = detected.get("encoding", "utf-8") or "utf-8"
    except ImportError:
        encoding = "utf-8"

    try:
        return content.decode(encoding, errors="replace")
    except Exception:
        return content.decode("utf-8", errors="replace")


def parse_csv(content: bytes) -> str:
    """Convert CSV to readable text."""
    try:
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return ""

        headers = rows[0]
        lines = [" | ".join(headers)]
        lines.append("-" * 40)
        for row in rows[1:]:
            lines.append(" | ".join(row))

        return "\n".join(lines)
    except Exception as e:
        logger.error("Failed to parse CSV: %s", e)
        return ""


def parse_document(filename: str, content: bytes, content_type: str = "") -> str:
    """Parse a document based on its filename extension or content type."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf" or "pdf" in content_type:
        return parse_pdf(content)
    elif ext == ".docx" or "wordprocessingml" in content_type:
        return parse_docx(content)
    elif ext == ".csv" or "csv" in content_type:
        return parse_csv(content)
    elif ext in (".txt", ".md", ".markdown", ".rst", ".log"):
        return parse_txt(content)
    else:
        # Try plain text as fallback
        return parse_txt(content)
